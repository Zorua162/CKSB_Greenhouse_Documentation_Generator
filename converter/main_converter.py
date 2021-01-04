# Module for editing docx files
import cv2
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX

# Module for date for auto last updated
from datetime import date

# Module for parsiing the config file
import yaml

# Warn if something goes wrong
import warnings

from pathlib import Path
import converter.images as images


CONFIG_PATH = "./biomes.yml/"
TEXTURES_PATH = "./textures/"
ICONS_PATH = "./icons/"
SAVE_PATH = "./GreenhouseDocumentation.docx"

COVERAGE_TYPES = ['water', 'lava', 'ice']

COLOUR_DICT = {
        '4':	'AA0000',
        'c':	'FF5555',
        '6':	'FFAA00',
        'e':	'FFFF55',
        '2':	'00AA00',
        'a':	'55FF55',
        'b':	'55FFFF',
        '3':	'00AAAA',
        '1':	'0000AA',
        '9':	'5555FF',
        'd':	'FF55FF',
        '5':	'AA00AA',
        'f':	'FFFFFF',
        '7':	'AAAAAA',
        '8':	'555555',
        '0':	'000000'
        }


def humanify(string):
    """
    Lower, capitalize and set underscores to spaces in given string
    """
    return string.lower().capitalize().replace('_', ' ')


def get_icon(icon):
    """
    Get the icon string from the biomes data
    """
    icon = icon.lower()
    icon_blockless = icon.replace('_block', '')
    icon_top = icon_blockless + '_top'
    to_search = [icon + '.png', icon_blockless + '.png', icon_top + '.png']
    textures = Path("./textures/")
    for search in to_search:
        searching = list(textures.glob(f"**/{search}"))
        if len(searching) != 0:
            return str(searching[0])
    print(f'Could not find the icon for {icon}')
    return None


def get_requirements(biome):
    """
    Turns contents: into a formatted string of block - number pairs for display
    """
    requirement = ''
    # Check biome has required blocks, avoids error if none required
    if "contents" in biome.keys():
        must_cont = biome["contents"].items()
        # For each required block add a line to the list for it
        content = [f"   {humanify(key)} - {value}" for key, value in must_cont]
        requirement += "\n".join(content)
    return requirement


def add_fluid_requirement(biome, paragraph):
    """
    Searches the config and adds fluid requirements if they are present
    """
    # First check that the biome requires a fluid, if so add heading for it
    if any((fld + 'coverage') in biome.keys() for fld in COVERAGE_TYPES):
        run = paragraph.add_run('\nFloor Coverage')
        run.bold = True
    # Check which of the fluids needs adding and format/ add it
    for fluid in COVERAGE_TYPES:
        fluid_key = fluid + 'coverage'
        if fluid_key in biome.keys():
            # If the fluid coverage is 0 then it means none is allowed
            if biome[fluid_key] > 0:
                paragraph.add_run(f'\n  {fluid.capitalize()} -> '
                                  f'{biome[fluid_key]}%')
            elif biome[fluid_key] == 0:
                paragraph.add_run(f'\n  No {fluid.capitalize()}')


def format_rgb(name):
    """
    format the rgb code given by the table in COLOUR_DICT into three codes
    colour: 6 digit code that can be split into r, g, b
    returns: r, g, b in the format: 0xXX, 0xXX, 0xXX where XX are colours
    """
    # Pull the mc colour code from the friendly name
    rgb = COLOUR_DICT[name[1:2]]

    # Format it to be three colours 0-255 rgb format
    r, g, b = (int(rgb[0:2], 16),
               int(rgb[2:4], 16),
               int(rgb[4:6], 16))
    return r, g, b


def create_doc():
    """
    Creates the Greenhouse Biomes Documentation
    """
    # Creates a blank template document
    document = Document()
    for style in document.styles:
        # Set for every style except list no
        if style.name != 'No List':
            style.font.name = 'Arial'
    # changing the page margins to be the same as the old marigins
    sections = document.sections
    for section in sections:
        margin = 1.27
        section.top_margin = Cm(margin)
        section.bottom_margin = Cm(margin)
        section.left_margin = Cm(margin)
        section.right_margin = Cm(margin)
    # key table taken from the original document
    table_key = {
        'Icon': 'What shows up in the in game inventory.',
        'Contents': 'Required blocks for the greenhouse to be valid.',
        'Plants': 'What plants can spawn if bonemeal is supplied to the '
        'greenhouse.',
        'Mobs': 'What mobs can spawn.',
        'Mob limit': 'Area of the greenhouse required to spawn a '
        'mob. For example if this is 9 then a greenhouse of area 9 will be '
        'able to spawn 1 mob.',
        'Floor Coverage': 'How much of the floor needs to be that '
        'kind of block. ',
        'Conversions': 'What blocks will convert to another block.',
        'Biome': 'What biome it will be inside of the greenhouse.'
    }

    # Add the title
    heading = document.add_heading('', 0)
    run = heading.add_run('Greenhouses Biomes Config')
    run.font.size = 14
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.bold = True
    run.underline = True

    # Not working warning:
    strikethrough = document.add_paragraph('')
    run = strikethrough.add_run('# strike through highlighted text '
                                '= not functioning')
    run.italics = True
    run.font.highlight_color = WD_COLOR_INDEX.YELLOW

    # Get the current date
    today = date.today().strftime("%d/%m/%Y")
    # Write this date to the file as the last updated date
    updated_paragraph = document.add_paragraph('')
    updated_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = updated_paragraph.add_run(f'Public Doc\nLast updated: {today}')
    run.italic = True
    run.font.size = Pt(8)

    # Add the key to the top of the document
    key_paragraph = document.add_paragraph('')
    for key, value in table_key.items():
        key = key_paragraph.add_run(f'{key}:')
        key.bold = True
        value = key_paragraph.add_run(f' {value}\n')

    # Next create the document

    # First parse all the data into a list of dictionary
    config_data = Path(CONFIG_PATH).read_text()
    biomes_data = yaml.load(config_data, Loader=yaml.SafeLoader)
    for biome in biomes_data["biomes"].values():
        # Initialise the paragraph
        biome_paragraph = document.add_paragraph('')
        # Add an empty run here for the icon to be added to later
        icon_img_run = biome_paragraph.add_run('')
        # Item 0 the friendly name
        if 'friendlyname' in biome:
            name = biome["friendlyname"]
            # format the friendly name depending on the colour codes
            if '&' in name:
                friendly_name = "\n" + name[2:] + ":"
                name_run = biome_paragraph.add_run(friendly_name)
                name_run.font.color.rgb = RGBColor(*format_rgb(name))
            else:
                friendly_name = "\n" + name + ":"
                name_run = biome_paragraph.add_run(friendly_name)

            name_run.bold = True
            name_run.font.size = Pt(13)

        # Cell 1 the icon
        # Row that contains the icon
        icon = biome['icon']
        icon_file = get_icon(icon)
        # now the image has been found add it through the paragraph and run
        icon_run = biome_paragraph.add_run("\nIcon: ")
        icon_run.bold = True
        biome_paragraph.add_run(humanify(icon))

        img = images.load_and_color(icon_file, 50, 150, 50)
        img_path = Path(f"{ICONS_PATH}/{icon.lower()}.png")
        img_path.parent.mkdir(parents=True, exist_ok=True)
        cv2.imwrite(str(img_path), img)

        icon_img_run.add_picture(str(img_path), width=Cm(1), height=Cm(1))

        # Cell 2 add the biome
        biome_run = biome_paragraph.add_run("\nBiome: ")
        biome_run.bold = True
        biome_run = biome_paragraph.add_run(humanify(biome['biome']))

        # Cell 3 add the requirements
        requirements = get_requirements(biome)
        requirements_run = biome_paragraph.add_run('\nContents:\n')
        requirements_run.bold = True
        requirements_run = biome_paragraph.add_run(requirements)

        # Add the Floor Coverage requirement
        add_fluid_requirement(biome, biome_paragraph)

        # Cell 4 add the Conversions
        if "conversions" in biome:
            conversions = ''
            # print(biome['conversions'])
            for input_block, conversion in biome['conversions'].items():
                # print(input_block, conversion)
                output = conversion.split(':')
                if len(output) == 2:
                    percentage, output_block = output
                    conversion = f'{humanify(input_block)} -> '
                    '{humanify(output_block)} : {percentage}% chance.'
                elif len(output) == 3:
                    percentage, output_block, adjacent = output
                    conversion = f'{humanify(input_block)} -> '
                    '{humanify(output_block)} : {percentage}% chance.'
                conversions += f'\n  ' + conversion
            conversions = biome_paragraph.add_run(conversions)

        # Cell 5 add the plants
        plants_string = ''
        if 'plants' in biome.keys():
            for plant, data in biome['plants'].items():
                chance, grows_on = data.split(':')
                plants_string += f'\n  {humanify(plant)} - {chance}% chance to'
                plants_string += f' grow on {humanify(grows_on)} .'
            run = biome_paragraph.add_run('\nPlants: ')
            run.bold = True
            biome_paragraph.add_run(plants_string)

        # Cell 6 add the mobs
        if 'mobs' in biome.keys():
            run = biome_paragraph.add_run('\nMobs:\n    ')
            run.bold = True
            for mob, data in biome['mobs'].items():
                chance, spawns_on = data.split(':')
                # Mob name is in italics so this must be in a seperate run
                run = biome_paragraph.add_run(humanify(mob))
                # Removed italics for consistency throughout the document
                # run.italic = True
                # now add the chance and block
                # some mobs can spawn in water so check if the block is water
                spawn_loc = 'on '
                if spawns_on == 'water':
                    spawn_loc = 'in'
                mob_string = f' - {chance}% {spawn_loc} {humanify(spawns_on)}'
                biome_paragraph.add_run(mob_string)

            if 'moblimit' in biome.keys():
                run = biome_paragraph.add_run('\nMoblimit: ')
                run.bold = True
                biome_paragraph.add_run(str(biome['moblimit']))

            else:
                warnings.warn(f'moblimit not set for {biome["biome"]}')

        # Cell 7 add the permissions
        if 'permission' in biome.keys():
            permission_string = '\n'
            if 'player.overworld' in biome['permission']:
                permission_string += '  This biome can only be made in ' \
                                     'the overworld'
            if 'player.nether' in biome['permission']:
                permission_string += '  This biome can only be made in ' \
                                     'the nether'
            if 'biome.nether' in biome['permission']:
                permission_string += 'This biome is possible to build in the' \
                                     ' overworld\nHowever it is exclusive ' \
                                     'to Donators and Trusted ranked players '
            run = biome_paragraph.add_run('\nPermissions: ')
            run.bold = True
            biome_paragraph.add_run(permission_string)

    # document.add_page_break()

    document.save(SAVE_PATH)


if __name__ == '__main__':
    create_doc()
