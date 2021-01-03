# Module for editing docx files
from colorsys import rgb_to_hls, hls_to_rgb
import cv2
import numpy as np
from skimage import io
from tqdm import tqdm


from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
# Module for date for auto last updated
from datetime import date
# Module for parsiing the config file
import yaml
# For getting icons
from os import listdir
# Path for creating icon path
from os.path import join
# Warn if something goes wrong
import warnings


def humanify(string):
    '''
    Lower, capitalize and set underscores to spaces in given string
    '''
    return string.lower().capitalize().replace('_', ' ')


def get_icon(icon):
    icon = icon.lower()
    icon_blockless = icon.replace('_block', '')
    icon_top = icon_blockless + '_top'
    # print(listdir('./block'))
    iconimg = icon + '.png'
    iconimg_blockless = icon_blockless + '.png'
    iconimg_top = icon_top + '.png'
    block_dir = listdir('./block')
    if iconimg in listdir('./item'):
        imagepath = join(r'./item', iconimg)
    elif iconimg in block_dir:
        imagepath = join(r'./block', iconimg)
    elif iconimg_blockless in block_dir:
        imagepath = join(r'./block', iconimg_blockless)
    elif iconimg_top in block_dir:
        imagepath = join(r'./block', iconimg_top)
    else:
        warnings.warn(f'Could not find the icon for this item')
        imagepath = None
    return imagepath


def get_requirements(biome):
    '''
    Searches the config and gets the block requirements for the biome
    '''
    # requirement = 'Blocks required:\n'
    requirement = ''
    if "contents" in biome.keys():
        content = ["  " + humanify(key) +
                   " - " + str(value) for key, value in biome["contents"].items()]
        requirement += "\n".join(content)
    return requirement


def add_fluid_requirement(biome, paragraph):
    '''
    Searches the config and adds fluid requirements if they are present
    '''
    coverage_types = ['water', 'lava', 'ice']
    if any((fl_type + 'coverage') in biome.keys() for fl_type in coverage_types):
        run = paragraph.add_run('\nFloor Coverage')
        run.bold = True

    for fluid in coverage_types:
        fluid_key = fluid + 'coverage'
        if fluid_key in biome.keys():
            if biome[fluid_key] > 0:
                paragraph.add_run(f'\n  {biome[fluid_key]}% {fluid.capitalize()}')
            elif biome[fluid_key] == 0:
                paragraph.add_run(f'\n  No {fluid.capitalize()}')


def fix_channels(img, fix_invert=True, force_trans=-1):
    """
    Forces images to be in RGBA format.
    """
    try:
        height, width, channels = img.shape
    except ValueError:
        height, width = img.shape
        channels = 1
    pixels = np.int8(img.reshape(width, -1, channels))
    oned = []
    if channels == 1:
        for w in range(width):
            for h in range(height):
                pixel = pixels[w, h] % 255
                if force_trans < 0:
                    force_trans = 255
                oned.extend([pixel, pixel, pixel, force_trans])
    else:
        for w in range(width):
            # oned = []
            for h in range(height):
                if channels == 4 and force_trans < 0:
                    pixel = pixels[w, h]
                    if fix_invert:
                        oned.extend([pixel[2] % 255, pixel[1] % 255, pixel[0] % 255, pixel[3] % 256])
                    else:
                        oned.extend([pixel[0] % 255, pixel[1] % 255, pixel[2] % 255, pixel[3] % 256])
                else:
                    pixel = pixels[w, h]
                    if force_trans < 0:
                        force_trans = 255
                    if fix_invert:
                        oned.extend([pixel[2] % 255, pixel[1] % 255, pixel[0] % 255, force_trans])
                    else:
                        oned.extend([pixel[0] % 255, pixel[1] % 255, pixel[2] % 255, force_trans])
        # last.append(oned)
    ar = np.array(oned)
    ar = ar.reshape(height, -1, 4)
    return ar


def map_sat(r1, g1, b1, r2, g2, b2):
    h1, l1, s1 = rgb_to_hls(r1 / 255, g1 / 255, b1 / 255)
    h2, l2, s2 = rgb_to_hls(r2 / 255, g2 / 255, b2 / 255)
    r3, g3, b3 = hls_to_rgb(h2, l1, (s1 + s2) / 2)
    return int(r3 * 255), int(g3 * 255), int(b3 * 255)


def color_image(img, r, g, b):
    pixels = fix_channels(io.imread(img))
    height, width, channels = img.shape
    img_rows = []
    img_stitched = None
    for h in range(height):
        row_image = None
        for w in range(width):
            p = pixels[h, w]
            a = p[3] % 256
            if a > 0:
                # If we don't copy, it actually gets modified.
                i = p.copy()
                new = map_sat(i[0], i[1], i[2], r, g, b)
                i[0] = new[0]
                i[1] = new[1]
                i[2] = new[2]
                i[3] = a
            else:
                i = p
            i = i.reshape((1, 1, 4))
            if row_image is None:
                row_image = i
            else:
                row_image = np.concatenate((row_image, i), axis=1)

        img_rows.append(row_image)
    for i in img_rows:
        if img_stitched is None:
            img_stitched = i
        else:
            img_stitched = np.concatenate((img_stitched, i), axis=0)

    return img


def create_doc():
    '''
    Creates the Greenhouse Biomes Documentation
    '''
    # Creates a blank template document
    document = Document()
    # for style in document.styles:
    #    print(style)

    # key table taken from the original document
    table_key = {
        'Icon': 'What shows up in the in game inventory.',
        'Contents': 'Required blocks for the greenhouse to be valid.',
        'Plants': 'What plants can spawn if bonemeal is supplied to the greenhouse.',
        'Mobs': 'What mobs can spawn.',
        'Mob limit': 'If the amount of mobs inside the greenhouse exceeds this, mobs ' +
                     'won’t spawn in the greenhouse.',
        'Floor Coverage': 'How much of the floor needs to be that kind of block. ',
        'Conversions': 'What blocks will convert to another block.',
        'Biome': 'What biome it will be inside of the greenhouse.'}

    # Assume that the config to use is ./biomes.yml
    # file_to_use = input('Enter the name of the biomes.yml to use >')

    # Add the title
    heading = document.add_heading('', 0)
    run = heading.add_run('Greenhouses Biomes Guide')
    run.bold = True
    run.italics = True

    # Not working warning:
    strikethrough = document.add_paragraph('')
    run = strikethrough.add_run('# strike through highlighted text ' +
                                '= not functioning')
    run.italics = True
    run.font.highlight_color = WD_COLOR_INDEX.YELLOW

    # Get the current date
    today = date.today().strftime("%d/%m/%Y")
    # Write this date to the file as the last updated date
    updated_paragraph = document.add_paragraph('')
    updated_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = updated_paragraph.add_run('Public Doc\nLast ' +
                                    f'updated: {today}')
    run.italic = True
    run.font.size = Pt(8)

    # Add the key to the top of the document
    key_paragraph = document.add_paragraph('')
    for key, value in table_key.items():
        key = key_paragraph.add_run(f'{key}:')
        key.bold = True
        value = key_paragraph.add_run(f' {value}\n')

    # Next create the document

    # First parse all the data into a list of dictionary
    with open("./biomes.yml") as file_object:
        biomes_data = yaml.load(file_object, Loader=yaml.SafeLoader)

    for biome in biomes_data["biomes"].values():
        # Initialise the paragraph
        biome_paragraph = document.add_paragraph('')
        # Add an empty run here for the icon to be added to later
        icon_img_run = biome_paragraph.add_run('')
        # Item 0 the friendly name
        if 'friendlyname' in biome:
            name = biome["friendlyname"]
            # format the friendly name depending on the colour codes
            if '&' in name:
                friendly_name = "\n" + name[2:][:len(name) - 5] + ":"
            else:
                friendly_name = "\n" + name + ":"

            name_run = biome_paragraph.add_run(friendly_name)
            name_run.bold = True
            name_run.font.size = Pt(13)

        # Cell 1 the icon
        # Row that contains the icon
        icon = biome['icon']
        icon_file = get_icon(icon)
        # now the image has been found add it through the paragraph and run
        icon_run = biome_paragraph.add_run("\nIcon: ")
        icon_run.bold = True
        biome_paragraph.add_run(humanify(icon))
        icon_img_run.add_picture(icon_file, width=Cm(1), height=Cm(1))

        # Cell 2 add the biome
        biome_run = biome_paragraph.add_run("\nBiome: ")
        biome_run.bold = True
        biome_run = biome_paragraph.add_run(humanify(biome['biome']))

        # Cell 3 add the requirements
        requirements = get_requirements(biome)
        requirements_run = biome_paragraph.add_run('\nRequirements:\n')
        requirements_run.bold = True
        requirements_run = biome_paragraph.add_run(requirements)

        # Add the Floor Coverage requirement
        add_fluid_requirement(biome, biome_paragraph)

        # Cell 4 add the Conversions
        if "conversions" in biome:
            conversions = ''
            # print(biome['conversions'])
            for input_block, conversion in biome['conversions'].items():
                # print(input_block, conversion)
                output = conversion.split(':')
                if len(output) == 2:
                    percentage, output_block = output
                    conversion = f'{humanify(input_block)} -> {humanify(output_block)} : {percentage}% chance.'
                elif len(output) == 3:
                    percentage, output_block, adjacent = output
                    conversion = f'{humanify(input_block)} -> {humanify(output_block)} : {percentage}% chance.'
                conversions += f'\n  ' + conversion
            conversions = biome_paragraph.add_run(conversions)

        # Cell 5 add the plants
        plants_string = ''
        if 'plants' in biome.keys():
            for plant, data in biome['plants'].items():
                chance, grows_on = data.split(':')
                plants_string += '\n  ' + humanify(plant) + ' will grow with a '
                plants_string += chance + '% on ' + humanify(grows_on)
            run = biome_paragraph.add_run('\nPlants: ')
            run.bold = True
            biome_paragraph.add_run(plants_string)

        # Cell 6 add the mobs
        if 'mobs' in biome.keys():
            mobs_string = ''
            for mob, data in biome['mobs'].items():
                chance, spawns_on = data.split(':')
                mobs_string += '\n  ' + humanify(mob) + ' will spawn on '
                mobs_string += humanify(spawns_on)
                run = biome_paragraph.add_run('\nMobs: ')
                run.bold = True
            biome_paragraph.add_run(mobs_string)
            if 'moblimit' in biome.keys():
                # mobs_string += f'\n The mob limit for this greenhouse is {biome["moblimit"]}'
                run = biome_paragraph.add_run('\nMoblimit: ')
                run.bold = True
                biome_paragraph.add_run(str(biome['moblimit']))

            else:
                warnings.warn(f'moblimit not set for {biome["biome"]}')

        # Cell 7 add the permissions
        if 'permission' in biome.keys():
            permission_string = '\n'
            if 'player.overworld' in biome['permission']:
                permission_string += "  This biome can only be made in the overworld"
            if 'player.nether' in biome['permission']:
                permission_string += "  This biome can only be made in the nether"
            if 'biome.nether' in biome['permission']:
                permission_string += '  This biome is possible to build in the overworld'
                permission_string += '\nHowever it is exclusive to Donators'
                permission_string += ' and Trusted ranked players'
            run = biome_paragraph.add_run('\nPermissions: ')
            run.bold = True
            biome_paragraph.add_run(permission_string)

    document.add_page_break()

    document.save('GreenhouseDocumentation.docx')


if __name__ == '__main__':
    create_doc()
